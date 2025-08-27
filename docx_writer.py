import base64
import io
import logging
import os

from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)


class DocxWriter:
    """Handles writing translated content to DOCX format while preserving layout."""
    
    def __init__(self, output_path: str = None):
        self.document = Document()
        self._setup_document_styles()
        self.output_path = output_path
        self.chunk_count = 0
        self.is_incremental_mode = output_path is not None
        
        # Track progress for incremental writing
        self.total_chunks = 0
        self.completed_chunks = 0
        
        if self.is_incremental_mode:
            logger.info(f"DocxWriter initialized in incremental mode for: {output_path}")
    
    def _setup_document_styles(self):
        """Setup document styles for different content types."""
        # Check python-docx version for compatibility
        try:
            import docx
            docx_version = getattr(docx, '__version__', '0.0.0')
            has_space_after = True
            logger.info(f"python-docx version: {docx_version}")
        except ImportError:
            docx_version = '0.0.0'
            has_space_after = False
            logger.warning("Could not determine python-docx version")
        
        # Define styles with their properties
        styles_config = [
            ('CustomHeader', {'font_name': 'Times New Roman', 'font_size': 16, 'bold': True, 'space_after': 12}),
            ('CustomTitle', {'font_name': 'Times New Roman', 'font_size': 14, 'bold': True, 'space_after': 10}),
            ('CustomNormal', {'font_name': 'Times New Roman', 'font_size': 12, 'bold': False, 'space_after': 6}),
            ('CustomFooter', {'font_name': 'Times New Roman', 'font_size': 9, 'bold': False, 'italic': True, 'space_after': 6}),
            ('CustomTableHeader', {'font_name': 'Times New Roman', 'font_size': 10, 'bold': True, 'space_after': 3})
        ]
        
        # Create all styles
        for style_name, config in styles_config:
            style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = config['font_name']
            style.font.size = Pt(config['font_size'])
            if config.get('bold'):
                style.font.bold = True
            if config.get('italic'):
                style.font.italic = True
            if has_space_after and hasattr(style, 'space_after'):
                style.space_after = Pt(config['space_after'])
    
    def write_chunk(self, translated_text: str = None, chunk_type: str = "text"):
        """Write a single translated chunk to the document and save immediately."""

        if translated_text is None:
            return

        try:
            # Add the translated content
            self._add_translated_chunk(translated_text, chunk_type)
            
            # Update progress
            self.completed_chunks += 1
            self._update_progress()
            
            # ✅ FIXED: Save after EVERY chunk to ensure immediate persistence
            if self.output_path:
                self._save_incremental_progress()
                
        except Exception as e:
            raise
    
    def _update_progress(self):
        """Update the progress paragraph in the document."""
        try:
            # Update the progress paragraph text
            progress_text = f"Translation Progress: {self.completed_chunks}/{self.total_chunks} chunks completed"
            if self.progress_paragraph and self.progress_paragraph.runs:
                # Clear existing text and add new text
                for run in self.progress_paragraph.runs:
                    run.text = ""
                self.progress_paragraph.runs[0].text = progress_text
            else:
                # Fallback: add a new progress paragraph
                self.progress_paragraph = self.document.add_paragraph(progress_text)
                self.progress_paragraph.style = self.document.styles['CustomNormal']
                
        except Exception as e:
            logger.warning(f"Could not update progress display: {e}")
    
    def _save_incremental_progress(self):
        """Save the document with current progress."""
        if not self.output_path:
            return
            
        try:
            self.document.save(self.output_path)
            completion_percent = (self.completed_chunks / self.total_chunks) * 100 if self.total_chunks > 0 else 0
            logger.info(f"✅ Chunk {self.completed_chunks} saved to file: {completion_percent:.1f}% complete ({self.completed_chunks}/{self.total_chunks} chunks)")
            
            # Verify file was actually written
            if os.path.exists(self.output_path):
                file_size = os.path.getsize(self.output_path)
                logger.info(f"📁 File verified: {self.output_path} ({file_size:,} bytes)")
            else:
                logger.warning(f"⚠️ File not found after save: {self.output_path}")
            
        except Exception as e:
            logger.error(f"❌ Failed to save incremental progress: {e}")
            # Don't raise exception here - we want translation to continue

    def _add_translated_chunk(self, translated_text: str, chunk_type: str = "text"):
        """Add a translated chunk to the document with appropriate styling."""
        if not translated_text.strip():
            return
        
        # Determine paragraph style based on chunk type
        if chunk_type == "header":
            style = 'CustomHeader'
        elif chunk_type == "title":
            style = 'CustomTitle'
        elif chunk_type == "footer":
            style = 'CustomFooter'
        elif chunk_type == "table_header":
            style = 'CustomTableHeader'
        else:
            style = 'CustomNormal'
        
        # Split text into paragraphs and write
        paragraphs = translated_text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                paragraph = self.document.add_paragraph(para_text)
                paragraph.style = self.document.styles[style]
                
                # Preserve bullet points and lists
                if para_text.startswith('•') or para_text.startswith('-'):
                    try:
                        paragraph.style = self.document.styles['List Bullet']
                    except KeyError:
                        # Fallback if List Bullet style doesn't exist
                        paragraph.style = self.document.styles['CustomNormal']
        
        # Add some spacing after the chunk
        self.document.add_paragraph("")  # Empty paragraph for spacing
    
    def add_metadata(self, title: str = "", author: str = "", subject: str = ""):
        """Add metadata to the document."""
        core_props = self.document.core_properties
        if title:
            core_props.title = title
        if author:
            core_props.author = author
        if subject:
            core_props.subject = subject

    def add_image_from_base64(self, base64_string: str, width: float = 4.0,
                              height: float = 3.0, caption: str = "", image_format: str = "PNG"):
        """
        Add an image to the document from a base64 encoded string.
        
        Args:
            base64_string (str): Base64 encoded image data
            width (float): Image width in inches (default: 4.0)
            height (float): Image height in inches (default: 3.0)
            caption (str): Optional caption below the image
            image_format (str): Image format (PNG, JPEG, etc.)
        
        Returns:
            bool: True if image was added successfully, False otherwise
        """
        try:
            # Decode base64 string
            image_data = base64.b64decode(base64_string)
            
            # Create PIL Image object to validate and potentially convert
            image = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if necessary (for JPEG compatibility)
            if image.mode in ('RGBA', 'LA', 'P'):
                # Create white background for transparent images
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = background
            
            # Save to BytesIO with specified format
            img_buffer = io.BytesIO()
            image.save(img_buffer, format=image_format)
            img_buffer.seek(0)
            
            # Add image to document
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run()
            
            # Add the image to the run
            run.add_picture(img_buffer, width=Inches(width), height=Inches(height))
            
            # Add caption if provided
            if caption:
                caption_para = self.document.add_paragraph(caption)
                caption_para.style = self.document.styles['CustomNormal']
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add some spacing after the image
            self.document.add_paragraph("")
            
            logger.info(f"✅ Image added successfully: {width}x{height} inches, format: {image_format}")
            return True
            
        except base64.binascii.Error as e:
            logger.error(f"❌ Invalid base64 string: {e}")
            return False
        except Exception as e:
            logger.error(f"❌ Failed to add image from base64: {e}")
            return False
